"""
Word document processor using python-docx.
"""

from pathlib import Path
from typing import Any

from loguru import logger

from judicaita.core.exceptions import DocumentParsingError, DocumentProcessingError
from judicaita.document_input.base import (
    DocumentContent,
    DocumentMetadata,
    DocumentProcessor,
)


class WordProcessor(DocumentProcessor):
    """Word document processor using python-docx."""

    def __init__(self, max_size_bytes: int = 50 * 1024 * 1024) -> None:
        """
        Initialize the Word processor.

        Args:
            max_size_bytes: Maximum allowed file size in bytes
        """
        self.max_size_bytes = max_size_bytes

    def supports(self, file_type: str) -> bool:
        """Check if Word format is supported."""
        return file_type.lower() in ["docx", ".docx", "doc", ".doc"]

    async def process(self, file_path: Path) -> DocumentContent:
        """
        Process a Word document and extract its content.

        Args:
            file_path: Path to the Word file

        Returns:
            DocumentContent: Extracted document content

        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            import docx

            # Validate file
            self._validate_file(file_path, self.max_size_bytes)

            logger.info(f"Processing Word document: {file_path}")

            # Load document
            doc = docx.Document(file_path)

            # Extract metadata
            metadata = await self._extract_metadata(file_path, doc)

            # Extract text content
            text_content: list[str] = []
            sections: list[dict[str, Any]] = []

            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    sections.append(
                        {
                            "index": para_idx,
                            "text": paragraph.text,
                            "style": paragraph.style.name,
                            "type": "paragraph",
                        }
                    )

            # Extract tables
            tables: list[dict[str, Any]] = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append(
                    {
                        "table_index": table_idx,
                        "data": table_data,
                        "rows": len(table.rows),
                        "cols": len(table.rows[0].cells) if table.rows else 0,
                    }
                )

            # Extract footnotes (if present in document properties)
            footnotes: list[str] = []
            # Note: python-docx has limited footnote support

            full_text = "\n\n".join(text_content)

            # Extract citations
            citations = self._extract_citations(full_text)

            return DocumentContent(
                text=full_text,
                metadata=metadata,
                sections=sections,
                tables=tables,
                footnotes=footnotes,
                citations=citations,
            )

        except ImportError as e:
            raise DocumentProcessingError(
                f"Required Word processing library not installed: {e}",
                error_code="MISSING_DEPENDENCY",
            )
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise DocumentParsingError(
                f"Failed to parse Word document: {e}",
                details={"file": str(file_path), "error": str(e)},
            )

    async def _extract_metadata(self, file_path: Path, doc: Any) -> DocumentMetadata:
        """Extract metadata from Word document."""
        try:
            core_props = doc.core_properties

            return DocumentMetadata(
                filename=file_path.name,
                file_type="docx",
                file_size_bytes=file_path.stat().st_size,
                page_count=None,  # python-docx doesn't provide page count
                author=core_props.author,
                title=core_props.title,
                subject=core_props.subject,
                created_date=str(core_props.created) if core_props.created else None,
                modified_date=str(core_props.modified) if core_props.modified else None,
                keywords=[k.strip() for k in (core_props.keywords or "").split(",") if k.strip()],
            )
        except Exception as e:
            logger.warning(f"Could not extract Word metadata: {e}")
            return DocumentMetadata(
                filename=file_path.name,
                file_type="docx",
                file_size_bytes=file_path.stat().st_size,
            )

    def _extract_citations(self, text: str) -> list[str]:
        """
        Extract legal citations from text using regex patterns.

        This is a simplified implementation. In production, use a specialized
        legal citation parser.
        """
        import re

        citations: list[str] = []

        # Pattern for US case citations
        case_pattern = r"\b\d+\s+[A-Z]\.[A-Za-z\.]+\s+\d+\b"
        citations.extend(re.findall(case_pattern, text))

        # Pattern for statute citations
        statute_pattern = r"\b\d+\s+U\.S\.C\.\s+§\s*\d+\b"
        citations.extend(re.findall(statute_pattern, text))

        return list(set(citations))  # Remove duplicates
